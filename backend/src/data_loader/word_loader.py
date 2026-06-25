from __future__ import annotations

from pathlib import Path
from typing import Iterator, Optional

from docx import Document as DocxDocument
from langchain_core.documents import Document

from backend.src.data_loader.base_loader import BaseDocumentLoader
from backend.src.data_loader.loader_registry import loader_registry


@loader_registry.register(extensions=[".docx"])
class DocxLoader(BaseDocumentLoader):
    def lazy_load(self) -> Iterator[Document]:
        doc = DocxDocument(str(self._file_path))

        full_text: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        if not full_text:
            yield Document(
                page_content="",
                metadata={"source_file": self._file_path.name, "warning": "empty_document"},
            )
            return

        for i, table in enumerate(doc.tables):
            table_lines: list[str] = []
            for row_idx, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                table_lines.append(" | ".join(cells))
            full_text.append(f"\n[Table {i + 1}]\n" + "\n".join(table_lines))

        combined = "\n\n".join(full_text)

        yield Document(
            page_content=combined,
            metadata={
                "source_file": self._file_path.name,
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            },
        )

    @classmethod
    def supports(cls, file_path: Path) -> bool:
        return file_path.suffix.lower() in (".docx",)


@loader_registry.register(extensions=[".doc"])
class DocLoader(BaseDocumentLoader):
    def lazy_load(self) -> Iterator[Document]:
        import warnings

        warnings.warn(
            "Legacy .doc format detected. Converting via antiword or LibreOffice is recommended."
        )
        try:
            import subprocess
            import tempfile

            ext_path = str(self._file_path)
            with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as tmp:
                tmp_path = tmp.name

            subprocess.run(["antiword", ext_path], stdout=open(tmp_path, "w"), check=True)
            text = Path(tmp_path).read_text(encoding="utf-8", errors="replace")
            Path(tmp_path).unlink(missing_ok=True)

            yield Document(
                page_content=text,
                metadata={"source_file": self._file_path.name, "converted_from": ".doc"},
            )
        except (FileNotFoundError, subprocess.CalledProcessError):
            yield Document(
                page_content=f"[Unsupported legacy .doc format. Please convert to .docx first: {self._file_path.name}]",
                metadata={"source_file": self._file_path.name, "error": "conversion_failed"},
            )

    @classmethod
    def supports(cls, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".doc"